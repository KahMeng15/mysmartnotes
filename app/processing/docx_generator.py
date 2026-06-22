"""
DOCX document generator for creating styled Word documents from structured content.
Supports headings, lists, images, and professional styling.
"""

import os
import logging
from typing import List, Optional
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

from app.processing.text_processor import ContentSegment, ContentType
from app.processing.image_extractor import ExtractedImage

logger = logging.getLogger(__name__)


class DocxGenerator:
    """Generates styled DOCX from structured content"""
    
    # Page size mapping (width, height) in cm
    PAGE_SIZES = {
        'A4': (21.0, 29.7),
        'LETTER': (21.59, 27.94),
        'LEGAL': (21.59, 35.56),
        'A3': (29.7, 42.0),
    }

    def __init__(
        self,
        resource_id: str,
        note_title: str,
        base_output_dir: str = "generated",
    ):
        self.resource_id = resource_id
        self.note_title = note_title
        self.base_output_dir = base_output_dir
        self.output_dir = os.path.join(base_output_dir, resource_id)

        # Create output directory
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)

        self.output_docx = os.path.join(self.output_dir, "OUTPUT.docx")

    def _set_cell_background(self, cell, fill_color):
        """Set background color for a table cell"""
        if not fill_color or fill_color.lower() == '#ffffff':
            return
        fill_color = fill_color.lstrip('#')
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_number_of_columns(self, section, cols):
        """Set number of columns for a section"""
        if cols <= 1:
            return
        sectPr = section._sectPr
        cols_el = sectPr.xpath('./w:cols')[0]
        cols_el.set(qn('w:num'), str(cols))
        cols_el.set(qn('w:space'), '720') # 0.5 inch gap default (720 twentieths of a point)

    def generate_docx(
        self,
        content_segments: List[ContentSegment],
        extracted_images: Optional[List[ExtractedImage]] = None,
        include_toc: bool = True,
        include_cover: bool = True,
        template_config: Optional[dict] = None,
        progress_callback=None,
    ) -> str:
        """
        Generate styled DOCX from content segments.
        Returns path to generated DOCX.
        """
        try:
            extracted_images = extracted_images or []
            self._template_config = template_config
            
            # Debug logging
            logger.info(f"[DOCX Export] Template config type: {type(template_config)}")
            logger.info(f"[DOCX Export] Template config value: {template_config}")
            
            doc = Document()

            if progress_callback:
                progress_callback("Preparing document", 10)

            # Configure default font and styles (applying template if provided)
            self._setup_styles(doc)

            # Cover page
            if include_cover:
                if progress_callback:
                    progress_callback("Creating cover page", 20)
                self._create_cover_page(doc)

            # Table of contents placeholder
            if include_toc:
                if progress_callback:
                    progress_callback("Building table of contents", 35)
                self._create_toc_placeholder(doc, content_segments)

            # Page size and orientation setup from template
            cols = 1
            if template_config and "page" in template_config:
                page_cfg = template_config["page"]
                size_name = page_cfg.get("size", "A4").upper()
                width_cm, height_cm = self.PAGE_SIZES.get(size_name, self.PAGE_SIZES['A4'])
                
                orientation = page_cfg.get("orientation", "portrait").lower()
                if orientation == "landscape":
                    width_cm, height_cm = height_cm, width_cm
                
                margins = page_cfg.get("margins", {"top": 25, "bottom": 25, "left": 19, "right": 19})
                cols = int(page_cfg.get("columns", 1))
                
                for section in doc.sections:
                    section.page_width = Cm(width_cm)
                    section.page_height = Cm(height_cm)
                    section.top_margin = Cm(margins.get("top", 25) / 10.0)
                    section.bottom_margin = Cm(margins.get("bottom", 25) / 10.0)
                    section.left_margin = Cm(margins.get("left", 19) / 10.0)
                    section.right_margin = Cm(margins.get("right", 19) / 10.0)
                    if orientation == "landscape":
                        section.orientation = WD_ORIENT.LANDSCAPE

            # Main content - Start a new section if columns > 1
            if cols > 1:
                from docx.enum.section import WD_SECTION
                new_section = doc.add_section(WD_SECTION.CONTINUOUS)
                self._set_number_of_columns(new_section, cols)
            
            if progress_callback:
                progress_callback("Rendering content", 50)
            self._create_content(doc, content_segments, extracted_images)

            # Footer with branding
            if progress_callback:
                progress_callback("Adding footer", 80)
            self._create_footer(doc)

            # Save
            if progress_callback:
                progress_callback("Saving document", 90)
            doc.save(self.output_docx)
            
            if progress_callback:
                progress_callback("Complete", 100)
            
            logger.info(f"Generated DOCX: {self.output_docx}")
            return self.output_docx

        except Exception as e:
            logger.error(f"Error generating DOCX: {e}", exc_info=True)
            raise

    def _set_font(self, font_obj, font_name: str):
        """
        Robustly set font name for a font object.
        Sets both the python-docx property and the underlying XML elements
        to ensure Word picks it up for all character sets (ascii, hAnsi, eastAsia, cs).
        """
        font_obj.name = font_name
        rFonts = font_obj._element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    def _setup_styles(self, doc: Document):
        """Configure default document styles, applying template if available"""
        tc = getattr(self, '_template_config', None) or {}
        el_cfg = tc.get("elements", {})
        spacing_cfg = tc.get("spacing", {})
        code_cfg = tc.get("code_block", {})
        
        font_family = tc.get("font_family", "Instrument Sans")
        
        # Helper to parse hex color
        def hex_to_rgb(hex_str):
            if not hex_str: return (0, 0, 0)
            hex_str = hex_str.lstrip('#')
            if len(hex_str) != 6: return (0, 0, 0)
            return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))
        
        # Normal style
        p_cfg = el_cfg.get("paragraph", {})
        style = doc.styles["Normal"]
        font = style.font
        self._set_font(font, font_family)
        font.size = Pt(p_cfg.get("font_size", 11))
        p_color = hex_to_rgb(p_cfg.get("text_color", "#000000"))
        font.color.rgb = RGBColor(*p_color)

        pf = style.paragraph_format
        pf.space_after = Pt(p_cfg.get("margin_bottom", spacing_cfg.get("paragraph_spacing", 8)))
        pf.space_before = Pt(p_cfg.get("margin_top", 0))
        pf.line_spacing = spacing_cfg.get("line_spacing", 1.15)

        # Handle Code style specially
        c_style = doc.styles["Code"] if "Code" in doc.styles else doc.styles.add_style('Code', 1)
        self._set_font(c_style.font, "Courier New")
        c_style.font.size = Pt(code_cfg.get("font_size", 9))
        c_color = hex_to_rgb(code_cfg.get("text_color", "#2c3e50"))
        c_style.font.color.rgb = RGBColor(*c_color)
        c_bg = code_cfg.get("background_color", "#f8f9fa")
        if c_bg and c_bg.lower() != '#ffffff':
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), c_bg.lstrip('#'))
            c_style.paragraph_format.element.get_or_add_pPr().append(shading_elm)
        c_style.paragraph_format.left_indent = Pt(15)
        c_style.paragraph_format.space_before = Pt(8)
        c_style.paragraph_format.space_after = Pt(8)

        # Heading styles from template or defaults
        heading_defaults = {
            "Heading 1": {"key": "h1", "size": 28, "color": "#1A1A2E", "bold": True, "space_before": 18, "space_after": 12},
            "Heading 2": {"key": "h2", "size": 24, "color": "#2C3E50", "bold": True, "space_before": 14, "space_after": 10},
            "Heading 3": {"key": "h3", "size": 20, "color": "#34495E", "bold": True, "space_before": 12, "space_after": 8},
            "Heading 4": {"key": "h4", "size": 16, "color": "#333333", "bold": True, "space_before": 10, "space_after": 6},
            "Title":     {"key": "note_title", "size": 28, "color": "#1A1A2E", "bold": True, "space_before": 24, "space_after": 12},
            "Subtitle":  {"key": "subject_name", "size": 18, "color": "#2C3E50", "bold": True, "space_before": 12, "space_after": 8},
            "List Bullet": {"key": "list", "size": 11, "color": "#2C3E50", "bold": False, "space_before": 0, "space_after": 4},
            "List Number": {"key": "list", "size": 11, "color": "#2C3E50", "bold": False, "space_before": 0, "space_after": 4},
        }
        
        # Add custom styles for GroupName if needed
        if "GroupName" not in doc.styles:
            style = doc.styles.add_style('GroupName', 1) # 1 = PARAGRAPH
            heading_defaults["GroupName"] = {"key": "group_name", "size": 12, "color": "#7f8c8d", "bold": False, "space_before": 4, "space_after": 8}

        # Setup Alignment mapping helper function
        def _get_docx_alignment(align_str):
            if align_str == 'center': return WD_ALIGN_PARAGRAPH.CENTER
            if align_str == 'right': return WD_ALIGN_PARAGRAPH.RIGHT
            if align_str == 'justify': return WD_ALIGN_PARAGRAPH.JUSTIFY
            return WD_ALIGN_PARAGRAPH.LEFT

        for style_name, defaults in heading_defaults.items():
            try:
                h_style = doc.styles[style_name]
                h_el = el_cfg.get(defaults["key"], {})
                self._set_font(h_style.font, font_family)
                h_style.font.size = Pt(h_el.get("font_size", defaults["size"]))
                h_color = hex_to_rgb(h_el.get("text_color", defaults["color"]))
                h_style.font.color.rgb = RGBColor(*h_color)
                h_style.font.bold = h_el.get("font_weight", "bold") == "bold"
                
                h_style.paragraph_format.space_before = Pt(h_el.get("margin_top", defaults["space_before"]))
                h_style.paragraph_format.space_after = Pt(h_el.get("margin_bottom", defaults["space_after"]))
                h_style.paragraph_format.alignment = _get_docx_alignment(h_el.get("alignment", "left"))
                h_style.paragraph_format.line_spacing = spacing_cfg.get("line_spacing", 1.15)
                
                # Apply background color (shading) if present
                bg_color = h_el.get("background_color")
                if bg_color and bg_color.lower() != '#ffffff':
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), bg_color.lstrip('#'))
                    h_style.paragraph_format.element.get_or_add_pPr().append(shading_elm)
                    
            except (KeyError, ValueError):
                pass

    def _create_cover_page(self, doc: Document):
        """Create a cover page using template settings"""
        tc = getattr(self, '_template_config', None) or {}
        cover_cfg = tc.get("cover_page", {})
        
        # Helper to parse hex color
        def hex_to_rgb(hex_str):
            hex_str = hex_str.lstrip('#')
            return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))
            
        # Vertical alignment via spacers
        v_align = cover_cfg.get("v_align", "middle")
        if v_align == "top":
            spacer_count = 1
        elif v_align == "middle":
            spacer_count = 8
        else: # bottom
            spacer_count = 16
            
        for _ in range(spacer_count):
            doc.add_paragraph("")

        # Horizontal alignment
        h_align_str = cover_cfg.get("h_align", "center")
        def _get_docx_alignment(align_str):
            if align_str == 'center': return WD_ALIGN_PARAGRAPH.CENTER
            if align_str == 'right': return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT
        h_align = _get_docx_alignment(h_align_str)

        # Title
        title_text = cover_cfg.get("h1_text") or self.note_title
        title_para = doc.add_paragraph()
        title_para.alignment = h_align
        run = title_para.add_run(title_text)
        run.font.size = Pt(cover_cfg.get("title_size", 36))
        t_color = hex_to_rgb(cover_cfg.get("title_color", "#1A1A2E"))
        run.font.color.rgb = RGBColor(*t_color)
        run.bold = True
        self._set_font(run.font, tc.get("font_family", "Instrument Sans"))

        # Subtitle
        subtitle_text = cover_cfg.get("h2_text") or "Extracted and Processed Content"
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = h_align
        run = subtitle_para.add_run(subtitle_text)
        run.font.size = Pt(cover_cfg.get("h2_size", 14))
        s_color = hex_to_rgb(cover_cfg.get("h2_color", "#7F8C8D"))
        run.font.color.rgb = RGBColor(*s_color)
        self._set_font(run.font, tc.get("font_family", "Instrument Sans"))

        # Date
        if cover_cfg.get("show_date", True):
            doc.add_paragraph("")
            date_para = doc.add_paragraph()
            date_para.alignment = h_align
            run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
            self._set_font(run.font, tc.get("font_family", "Instrument Sans"))

        # Page break after cover
        doc.add_page_break()

    def _create_toc_placeholder(self, doc: Document, content_segments: List[ContentSegment]):
        """Create a table of contents section"""
        tc = getattr(self, '_template_config', None) or {}
        font_family = tc.get("font_family", "Instrument Sans")
        
        toc_heading = doc.add_paragraph()
        run = toc_heading.add_run("Table of Contents")
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        self._set_font(run.font, font_family)

        # List headings from content
        for segment in content_segments:
            if segment.content_type in [ContentType.H1, ContentType.H2]:
                indent = "" if segment.content_type == ContentType.H1 else "    "
                content_text = segment.content[:60]
                if len(segment.content) > 60:
                    content_text += "..."
                entry = doc.add_paragraph(f"{indent}• {content_text}")
                entry.paragraph_format.space_after = Pt(4)
                for run in entry.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                    self._set_font(run.font, font_family)
                if segment.content_type == ContentType.H2:
                    entry.paragraph_format.left_indent = Pt(20)

        doc.add_page_break()

    def _prepare_content(self, segment: ContentSegment, style_name: str) -> str:
        """
        Prepare segment content:
        - Strip and replace list prefix chars with proper bullet / number
        """
        content = segment.content
        tc = getattr(self, '_template_config', None) or {}
        list_cfg = tc.get("list_config", {})
        custom_bullet = list_cfg.get("custom_bullet", "•")
        num_format = list_cfg.get("number_format", "1.")

        if style_name == "List Bullet":
            # Strip leading markdown list markers (-, *, +, •) and replace with custom_bullet
            content = re.sub(r'^\s*[-*+•]\s+', '', content)
            content = f'{custom_bullet}  {content}'
        elif style_name == "List Number":
            # Strip leading "1." / "2." etc. and replace with the configured format
            m = re.match(r'^\s*(\d+)\.\s+', content)
            if m:
                num = m.group(1)
                content = re.sub(r'^\s*\d+\.\s+', '', content)
                
                # Apply number format
                prefix = f"{num}."
                if num_format == "1)": prefix = f"{num})"
                elif num_format == "1]": prefix = f"{num}]"
                
                content = f'{prefix}  {content}'

        return content

    def _create_content(
        self,
        doc: Document,
        content_segments: List[ContentSegment],
        extracted_images: List[ExtractedImage],
    ):
        """Create main content using template styles"""
        tc = getattr(self, '_template_config', None) or {}
        # Image lookup by page
        images_by_page = {}
        for img in extracted_images:
            if img.page_number not in images_by_page:
                images_by_page[img.page_number] = []
            images_by_page[img.page_number].append(img)

        i = 0
        while i < len(content_segments):
            segment = content_segments[i]
            
            # Check if this is a table row - collect all consecutive table rows
            if segment.content_type == ContentType.TABLE_ROW:
                table_segments = []
                while i < len(content_segments) and content_segments[i].content_type == ContentType.TABLE_ROW:
                    table_segments.append(content_segments[i])
                    i += 1
                self._render_table(doc, table_segments)
                continue

            if segment.content_type == ContentType.NOTE_TITLE:
                doc.add_paragraph(segment.content, style="Title")
            elif segment.content_type == ContentType.SUBJECT_NAME:
                doc.add_paragraph(segment.content, style="Subtitle")
            elif segment.content_type == ContentType.GROUP_NAME:
                doc.add_paragraph(segment.content, style="GroupName")
            elif segment.content_type == ContentType.H1:
                doc.add_heading(segment.content, level=1)
            elif segment.content_type == ContentType.H2:
                doc.add_heading(segment.content, level=2)
            elif segment.content_type == ContentType.H3:
                doc.add_heading(segment.content, level=3)
            elif segment.content_type == ContentType.LIST:
                content = self._prepare_content(segment, "List Bullet")
                doc.add_paragraph(content, style="List Bullet")
            elif segment.content_type == ContentType.ORDERED_LIST:
                content = self._prepare_content(segment, "List Number")
                doc.add_paragraph(content, style="List Number")
            elif segment.content_type == ContentType.CODE:
                para = doc.add_paragraph()
                run = para.add_run(segment.content)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                para.paragraph_format.left_indent = Pt(15)
            else:
                # Body text
                doc.add_paragraph(segment.content)

            # Add images for this page
            if segment.page_number in images_by_page:
                for img in images_by_page[segment.page_number]:
                    self._add_image(doc, img)
                del images_by_page[segment.page_number]
            
            i += 1

    def _render_table(self, doc: Document, table_segments: List[ContentSegment]):
        """Render a table in DOCX using template settings"""
        if not table_segments: return
        
        # Parse rows
        rows_data = []
        for s in table_segments:
            # Simple markdown table parser
            cells = [c.strip() for c in s.content.split("|") if c.strip() and not c.strip().startswith("---")]
            if cells: rows_data.append(cells)
        if not rows_data: return

        tc = getattr(self, '_template_config', None) or {}
        tbl_cfg = tc.get("table", {})
        
        # Table settings
        header_fs = tbl_cfg.get("header_font_size", 10)
        body_fs = tbl_cfg.get("body_font_size", 9)
        align_str = tbl_cfg.get("alignment", "left").lower()
        
        hdr_bg = tbl_cfg.get("header_bg_color", "#593f8f")
        hdr_fg = tbl_cfg.get("header_text_color", "#ffffff")
        odd_bg = tbl_cfg.get("odd_row_color", "#ffffff")
        even_bg = tbl_cfg.get("even_row_color", "#f5f4fb")
        
        def _get_docx_alignment(a_str):
            if a_str == 'center': return WD_ALIGN_PARAGRAPH.CENTER
            if a_str == 'right': return WD_ALIGN_PARAGRAPH.RIGHT
            return WD_ALIGN_PARAGRAPH.LEFT
        h_align = _get_docx_alignment(align_str)

        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        table.style = 'Table Grid'
        
        # Helper to parse hex color to RGBColor
        def hex_to_rgb_obj(hex_str):
            if not hex_str: return RGBColor(0, 0, 0)
            hex_str = hex_str.lstrip('#')
            return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

        for r_idx, row in enumerate(rows_data):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= len(table.columns): continue
                cell = table.cell(r_idx, c_idx)
                
                # Background color
                if r_idx == 0:
                    self._set_cell_background(cell, hdr_bg)
                elif r_idx % 2 == 0: # Even row (0-indexed, so 2nd row of data is index 1, which is odd...)
                    # Wait, r_idx 0 is header. 
                    # Row 1 (data row 1) is odd.
                    # Row 2 (data row 2) is even.
                    self._set_cell_background(cell, even_bg)
                else: # Odd row
                    self._set_cell_background(cell, odd_bg)

                # Set text and formatting
                para = cell.paragraphs[0]
                para.alignment = h_align
                run = para.add_run(cell_text)
                run.font.size = Pt(header_fs if r_idx == 0 else body_fs)
                self._set_font(run.font, tc.get("font_family", "Instrument Sans"))
                
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = hex_to_rgb_obj(hdr_fg)
                
        doc.add_paragraph("")

    def _add_image(self, doc: Document, image_obj: ExtractedImage):
        """Add image to document"""
        try:
            if not image_obj.file_path or not os.path.exists(image_obj.file_path):
                logger.warning(f"Image file not found: {image_obj.file_path}")
                return

            # Caption before image
            if image_obj.caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption_para.add_run(image_obj.caption)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
                run.italic = True

            # Add image (max 5 inches wide)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_obj.file_path, width=Inches(5))

        except Exception as e:
            logger.error(f"Error adding image {image_obj.filename}: {e}")

    def _create_footer(self, doc: Document):
        """Add branded footer to all sections using template settings"""
        tc = getattr(self, '_template_config', None) or {}
        footer_cfg = tc.get("footer", {})
        
        custom_text = footer_cfg.get("custom_text", "")
        parts = ["Generated by mysmartnotes.vercel.app | Create notes and study smart!"]
        if custom_text: parts.append(custom_text)
        footer_text = " | ".join(parts)

        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = ""
            
            # Note: Automatic page numbering in DOCX footer is complex via python-docx.
            # We will just use the text for now.
            run = footer_para.add_run(footer_text)
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
            self._set_font(run.font, tc.get("font_family", "Instrument Sans"))
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
